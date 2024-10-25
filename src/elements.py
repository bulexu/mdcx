import re
import docx
import requests
import PIL.Image
from io import BytesIO
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .context import Context
from .utils import _add_link, _is_bib, _level_info
from copy import copy

STYLE_CODE = "Code"

class Heading:
    """Heading section inside document"""

    def __init__(self, text: str, level: int) -> None:
        self.text = text
        self.level = level

    def _md(line: str):
        # Parse number of # for level
        level = 0
        while line[level] == "#":
            level += 1
        # Get and clean text
        text = line[level:].lstrip()
        return Heading(text, level)

    def _docx(self, docx_doc: docx.Document):
        # Page break for bibliography
        if _is_bib(self.text):
            docx_doc.add_page_break()
        # Add heading
        docx_para = docx_doc.add_heading(self.text, self.level)
        # TODO: bookmarks
        # print(docx_para._element.xml)
        # docx_para.insert(0, etree.XML("<hi />"))



class Run:
    """Run of text with styling located inside a paragraph"""

    def __init__(self, ctx: Context, text: str, link=None, image=None):
        # Check that run is a string; python doesn't have strong typing sadly
        if type(text) != str:
            raise Exception("Make sure this run is a string, this is a common mistake")
        # Create tuns
        self.ctx = copy(ctx)
        self.text = text
        self.link = link
        self.image = image

    def _docx(self, docx_para):
        if self.image:
            url, alt_text, title = self.image
            img_data = None

            # 尝试获取图片数据
            if url.startswith(('http://', 'https://')):
                try:
                    response = requests.get(url, timeout=10)
                    response.raise_for_status()
                    img_data = BytesIO(response.content)
                except Exception as e:
                    print(f"无法下载图片 {url}: {e}")
                    docx_para.add_run(f"[图片: {url} 下载失败]")
            else:
                img_path = self.ctx.link_to(url)
                if img_path.exists():
                    img_data = str(img_path)
                else:
                    print(f"图片文件不存在: {url}")
                    docx_para.add_run(f"[图片: {url} 文件不存在]")

            if img_data:
                try:
                    # 获取图片尺寸
                    img = PIL.Image.open(img_data if isinstance(img_data, BytesIO) else img_data)
                    width, height = img.size
                    
                    # 插入图片
                    if height > width:
                        docx_para.add_run().add_picture(img_data, height=Cm(10))
                    else:
                        docx_para.add_run().add_picture(img_data, width=Cm(12))
                    
                    # 如果有标题,添加图片说明
                    if title:
                        docx_para.add_run().add_break()
                        docx_para.add_run(f"图 {self.ctx.figures} - {title}")
                        self.ctx.figures += 1
                except Exception as e:
                    print(f"无法插入图片 {url}: {e}")
                    docx_para.add_run(f"[图片: {url}]")
            else:
                # 如果无法获取图片,将链接作为文本输出
                docx_para.add_run(f"[图片: {url}]")
        elif self.link:
            link, external = self.link
            return _add_link(docx_para, link, self.text, external)
        else:
            # Add plain run text
            docx_run = docx_para.add_run(self.text)
            # Add relevant styles
            if self.ctx.bold:
                docx_run.bold = True
            if self.ctx.italic:
                docx_run.italic = True
            if self.ctx.underline:
                docx_run.underline = True
            if self.ctx.strikethrough:
                docx_run.strikethrough = True
            return docx_run


class Paragraph:
    """Paragraph consisting of many runs of text"""

    def __init__(self, ctx: Context, runs: list = []):
        self.ctx = ctx
        self.runs = runs

    def append(self, run: Run):
        """Appends new run to paragraph"""
        self.runs.append(run)

    @staticmethod
    def _md(ctx: Context, line: str):
        # Metadata
        runs = []
        ind = 0
        flipflop = False
        buf = ""

        # Go through each character
        while ind < len(line):
            # Flipflops
            if flipflop:
                buf += line[ind]
                ind += 1
                flipflop = False
            # Backslash for flipflop
            elif line[ind] == "\\":
                flipflop = True
            # Bold/italics
            elif line[ind] == "*":
                # Clear buf
                runs.append(Run(ctx, buf))
                buf = ""
                # Parse
                ind += _run_ib(ctx, line[ind:])
            # Cheeky link
            elif line[ind] == "<" and ">" in line[ind:]:
                # Clear buf
                runs.append(Run(ctx, buf))
                buf = ""
                # Parse
                res = _run_cheeky(ctx, line[ind:])
                ind += res[0]
                runs.append(res[1])
            # Misc
            else:
                # Find instances of link or image
                match = re.search(
                    r"^(!?)\[(.+?)\]\((.+?)(\s+\"(.+?)\")?\)",
                    line[ind:],
                )

                # Link or Image
                if match:
                    # Finish existing buffer and skip link/image
                    runs.append(Run(ctx, buf))
                    buf = ""
                    ind += len(match.group(0))
                    
                    is_image = bool(match.group(1))
                    text = match.group(2)
                    link = match.group(3)
                    title = match.group(5)

                    if is_image:
                        # Image
                        runs.append(Run(ctx, "", image=(link, text, title)))
                    else:
                        # Link
                        if link.startswith("#"):
                            # Internal link
                            runs.append(Run(ctx, text, link=(link[1:], False)))
                        else:
                            # External link
                            runs.append(Run(ctx, text, link=(link, True)))

                # Normal character
                else:
                    buf += line[ind]
                    ind += 1

        # Create paragraph and return
        runs.append(Run(ctx, buf))
        return Paragraph(ctx, runs)

    def _docx(self, docx_doc: docx.Document) -> docx.text.paragraph.Paragraph:
        # Add empty paragraph
        docx_para = docx_doc.add_paragraph()
        # Make no-spaced if defined
        if self.ctx.no_spacing():
            docx_para.style = "No Spacing"
        # Add runs to paragraph
        for run in self.runs:
            run._docx(docx_para)
        return docx_para


class Codeblock:
    """Codeblock containing language and monospaced code"""

    def __init__(self, lines: list, lang: str = None, heading_after: bool = False):
        self.lines = lines
        self.lang = lang  # TODO: use somewhere in docx
        self.heading_after = heading_after

    @staticmethod
    def _md(lines: list) -> tuple:
        # Get language after ``` designator
        lang = (
            lines[0].lstrip()[3:].lstrip()
        )  # first `lstrip()` used in document parsing
        lang = lang if lang != "" else None

        # Read lines
        heading_after = False
        code = []
        for ind, line in enumerate(lines[1:]):
            if line.lstrip() == "```":
                # Check if there's a heading afterwards
                if len(lines[1:]) - 1 > ind and lines[ind + 2].lstrip().startswith("#"):
                    heading_after = True
                # Stop codeblock
                break
            else:
                code.append(line)

        # Get skip
        skip = len(code) + 1
        return (Codeblock(code, lang, heading_after), skip)

    def _docx(self, docx_doc: docx.Document):
        # Calculate justification for lines
        just = len(str(len(self.lines)))
        # Add lines
        for ind, line in enumerate(self.lines):
            # Figure out line number
            num = str(ind + 1).rjust(just)
            # Add new paragraph with code style
            docx_para = docx_doc.add_paragraph()
            docx_para.style = STYLE_CODE
            # Add line number with italics
            docx_run = docx_para.add_run(num)
            docx_run.font.italic = True
            # Add actual code
            docx_para.add_run(" " + line)

        # Add small codeblock line for formatting if there's not a heading afterwards
        if not self.heading_after:
            docx_para = docx_doc.add_paragraph()
            docx_para.style = STYLE_CODE


class Quote(Paragraph):
    """Quote of something in it's own style"""

    @staticmethod
    def _md(ctx: Context, line: str):
        # Level info
        level, line = _level_info(line)
        # Clean line from `>` starter
        line = line[1:].lstrip()
        # Parse via inheritance and convert
        para = super(Quote, Quote)._md(
            ctx, line
        )  # BODGE: python doesn't like staticmethod and inheritance
        quote = Quote(ctx, para.runs)
        # Set levelling
        quote.level = level
        return quote

    def _docx(self, docx_doc: docx.Document) -> docx.text.paragraph.Paragraph:
        # Get inherited generated paragraph
        para = super()._docx(docx_doc)
        # Reset to quote styling
        para.style = "Quote"
        INDENT = 0.75
        para.paragraph_format.left_indent = Cm(INDENT * self.level + 1)
        para.paragraph_format.right_indent = Cm(INDENT)
        return para


class PointBullet(Paragraph):
    """Bullet point with content inside of it"""

    @staticmethod
    def _md(ctx: Context, line: str):
        # Level info
        level, line = _level_info(line)
        # Clean line from `-` starter
        line = line[1:].lstrip()
        # Parse via inheritance and convert
        para = super(PointBullet, PointBullet)._md(
            ctx, line
        )  # BODGE: python doesn't like staticmethod and inheritance
        bullet = PointBullet(ctx, para.runs)
        # Set levelling
        bullet.level = level
        return bullet

    def _docx(self, docx_doc: docx.Document) -> docx.text.paragraph.Paragraph:
        # Get inherited generated paragraph
        docx_para = super()._docx(docx_doc)
        # Set bullet style according to level
        docx_para.style = (
            "List Bullet" if self.level == 0 else f"List Bullet {self.level+1}"
        )
        return docx_para


class PointNumbered(Paragraph):
    """Numbered point with content inside of it"""

    @staticmethod
    def _md(ctx: Context, line: str):
        # Level info
        level, line = _level_info(line)
        # Get number and clean
        splitted = line.split(".", 1)
        num = int(splitted[0])
        line = splitted[1].lstrip()
        # Parse via inheritance and convert
        para = super(PointNumbered, PointNumbered)._md(
            ctx, line
        )  # BODGE: python doesn't like staticmethod and inheritance
        numbered = PointNumbered(ctx, para.runs)
        # Set info
        numbered.level = level
        numbered.num = num
        return numbered

    def _docx(self, docx_doc: docx.Document) -> docx.text.paragraph.Paragraph:
        # TODO: use something like "start at self.num" so markdown starting at like `20.` can be used, it fucks up otherwise
        # Get inherited generated paragraph
        docx_para = super()._docx(docx_doc)
        # Set bullet style according to level
        docx_para.style = (
            "List Number" if self.level == 0 else f"List Number {self.level+1}"
        )
        return docx_para


class Image:
    """Image with some optional caption text"""

    def __init__(self, ctx: Context, link: str, caption: Paragraph = None) -> None:
        # Get and check image link
        real_link = ctx.link_to(link)
        if not real_link.exists():
            raise Exception(f"Image linked to as {link} does not exist")

        # Set other values
        self.ctx = ctx
        self.link = real_link
        self.safe_link = str(real_link.absolute())
        self.caption = caption

    @staticmethod
    def _md(ctx: Context, matched: str):
        splitted = matched.split("](")
        caption = splitted[0][2:].strip()
        if caption != "":
            ctx.figures += 1
            caption = Paragraph._md(ctx, f"Figure {ctx.figures} - {caption}")
        else:
            caption = None
        link = splitted[1][:-1].strip()
        return Image(copy(ctx), link, caption)

    def _docx(self, docx_doc: docx.Document) -> list[docx.text.paragraph.Paragraph]:
        # Get image width/heigth
        img = PIL.Image.open(self.link)
        width, height = (img.width, img.height)

        # Insert image
        docx_para_image = docx_doc.add_paragraph()
        docx_run = docx_para_image.add_run()
        try:
            # Width/height adjustment so it won't fall off the page
            if height > width:
                docx_run.add_picture(self.safe_link, height=Cm(10))
            else:
                docx_run.add_picture(self.safe_link, width=Cm(12))
        except Exception as e:
            raise Exception(f"Failed to add image {self.link} to document ({e})")

        # Add caption
        if self.caption:
            docx_para_caption = self.caption._docx(docx_doc)
            docx_para_caption.style = "Caption"

            return [docx_para_image, docx_para_caption]
        return [docx_para_image]


class Table:
    def __init__(self, rows):
        self.rows = rows

    @staticmethod
    def _md(lines):
        rows = []
        skip = 0
        for line in lines:
            if not line.strip().startswith("|"):
                break
            cells = [cell.strip() for cell in line.split("|")[1:-1]]
            if all(cell.replace("-", "").replace(":", "") == "" for cell in cells):
                skip += 1
                continue
            rows.append(cells)
            skip += 1
        return Table(rows), skip

    def _docx(self, docx_doc):
        table = docx_doc.add_table(rows=len(self.rows), cols=len(self.rows[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row in enumerate(self.rows):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell
        
        # 设置单元格边框
        self.set_cell_border(table)

    @staticmethod
    def set_cell_border(table):
        tbl = table._tbl
        for cell in tbl.iter_tcs():
            tcPr = cell.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                edge = OxmlElement(f'w:{border}')
                edge.set(qn('w:val'), 'single')
                edge.set(qn('w:sz'), '4')
                edge.set(qn('w:space'), '0')
                edge.set(qn('w:color'), 'auto')
                tcBorders.append(edge)
            tcPr.append(tcBorders)


def _run_cheeky(ctx: Context, line: str) -> tuple:
    """Run parsing for cheeky links (the <> links)"""

    # Metadata
    link = ""
    flipflop = False

    # Go through each character
    for c in line[1:]:
        # Flipflop
        if flipflop:
            flipflop = False
            link += c
        # Backslash for flipflop
        elif c == "\\":
            flipflop = True
        # End of link
        elif c == ">":
            break
        # Character in link
        else:
            link += c

    # Construct new run
    run = Run(ctx, link, link=(link, True))

    # Return ind and link
    return len(link) + 2, run

def _run_ib(ctx: Context, line: str) -> int:
    """Run parsing for italics and bold"""

    # Get star count
    stars = len(line) - len(line.lstrip("*"))

    # Italics for non-even
    if stars % 2 == 1:
        ctx.flip_italic()

    # Bold if theres more than one, coexists with italics
    if stars > 1:
        ctx.flip_bold()

    # Add star count to index
    return stars