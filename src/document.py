import re
import docx
from copy import copy
from .elements import Paragraph, Heading, Run, Codeblock, Quote, PointBullet, Image, Table, PointNumbered
from .context import Context
from .styles import Style
from pathlib import Path
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from .utils import _style_title_border, _rm_toc


STYLE_CODE = "Code"

class Document:
    """High-level document abstractions for conversion"""

    def __init__(self, md: str, path: Path, style: Style = Style.andy()):
        # Components
        self.elements = []
        self.title = None
        self.subtitle = None
        self.ctx = Context(path.parent)
        self.style = style

        # Remove toc and clear up lines
        lines_raw = _rm_toc(md)
        lines = []
        for line in lines_raw:
            lines.append(line.rstrip())

        # Metadata
        if len(lines) > 1 and lines[0] == "---":
            # Go over lines in metadata
            skip = 0
            for ind, line in enumerate(lines[1:]):
                # Stop metadata if it's ended
                if line == "---":
                    skip = ind + 1
                    break
                # Split at `:` token
                splitted = line.split(":", 1)
                # Go to next line if its invalid
                if len(splitted) != 2:
                    continue
                # Clean left and right sections
                left = splitted[0].lstrip().lower()
                right = splitted[1].lstrip()
                # Match left section
                if left == "title":
                    self.title = right
                elif left == "subtitle":
                    self.subtitle = right
            # Skip to end of metadata if there was an open and close tag
            if skip != 0:
                lines = lines[1 + skip :]

        # Parse through lines
        while self.ctx.line < len(lines):
            # 获取当前行
            line = lines[self.ctx.line]
            stripped = line.lstrip()
            # 检查行首
            if stripped.startswith("<!--"):
                # 注释
                self.ctx.next_line()
                continue
            if stripped.startswith("#"):
                # 标题
                heading = Heading._md(stripped)
                self.elements.append(heading)
                self.ctx.heading = heading
            elif stripped.startswith("```"):
                # 代块
                codeblock, skip = Codeblock._md(lines[self.ctx.line :])
                self.ctx.line += skip
                self.elements.append(codeblock)
            elif stripped.startswith(">"):
                # 引用
                self.elements.append(Quote._md(copy(self.ctx), line))
            elif stripped.startswith("-"):
                # 无序列表
                self.elements.append(PointBullet._md(copy(self.ctx), line))
            elif match := re.search(
                r"^!\[.*\]\(.+\)",
                line,
            ):
                # 图片
                self.elements.append(Image._md(self.ctx, match.group(0)))
            elif stripped.startswith("|") or stripped.startswith("+-"):
                # 表格
                table, skip = Table._md(lines[self.ctx.line:])
                self.ctx.line += skip
                self.elements.append(table)
            # Check others
            else:
                # Numbered point
                try:
                    if "." not in stripped:
                        raise Exception()  # TODO: better error
                    int(stripped.split(".", 1)[0])
                    self.elements.append(PointNumbered._md(copy(self.ctx), line))
                # Paragraph
                except:
                    if (
                        # Non-sensitive typical empty lines
                        (not self.ctx.no_spacing() and line == "")
                        # Sensitive but last line was title
                        or (
                            self.ctx.no_spacing()
                            and lines[self.ctx.line - 1].lstrip().startswith("#")
                        )
                        # Sensitive but next line is title
                        or (
                            self.ctx.no_spacing()
                            and len(lines) > self.ctx.line + 1
                            and lines[self.ctx.line + 1].lstrip().startswith("#")
                        )
                    ):
                        # Skip empty line
                        self.ctx.next_line()
                        continue
                    self.elements.append(Paragraph._md(copy(self.ctx), stripped))

            # Move to next line
            self.ctx.next_line()

    def save(self, path: Path):
        """Saves document to `path` provided"""
        # Create docx file
        docx_doc = docx.Document()

        # New styles
        style_codeblock = docx_doc.styles.add_style(STYLE_CODE, WD_STYLE_TYPE.PARAGRAPH)

        # Add title/subtitle
        if self.title or self.subtitle:
            # Create empty lines before title
            for _ in range(4):
                para = Paragraph(copy(self.ctx), [Run(copy(self.ctx), "")])
                para._docx(docx_doc)

            # Add title
            if self.title:
                docx_para = docx_doc.add_heading(self.title, 0)
            # Add subtitle
            if self.subtitle:
                docx_para = Paragraph(
                    copy(self.ctx), [Run(copy(self.ctx), self.subtitle)]
                )._docx(docx_doc)
                docx_para.style = "Subtitle"

            # Page break
            docx_para = docx_doc.add_paragraph()
            docx_run = docx_para.add_run()
            docx_run.add_break(WD_BREAK.PAGE)

        # Add elements
        for element in self.elements:
            element._docx(docx_doc)

        # Replace all fonts with body font by default
        for style in docx_doc.styles:
            if hasattr(style, "font"):
                style.font.name = self.style.font_body

        # Styling for title
        style_title = docx_doc.styles["Title"]
        _style_title_border(style_title)
        style_title.font.name = self.style.font_heading
        style_title.font.size = Pt(26)
        if not self.style.heading_blue:
            style_title.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style_title.paragraph_format.space_after = Pt(3)
        style_title.paragraph_format.alignment = 1

        # Styling for subtitle
        style_subtitle = docx_doc.styles["Subtitle"]
        style_subtitle.font.name = self.style.font_heading
        style_subtitle.font.size = Pt(14)
        if not self.style.heading_blue:
            style_subtitle.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style_subtitle.font.italic = False
        style_subtitle.paragraph_format.alignment = 1

        # Styling for headings
        for h in range(1, 9):
            style_heading = docx_doc.styles[f"Heading {h}"]
            style_heading.font.name = self.style.font_heading
            style_heading.font.bold = self.style.heading_bold
            if not self.style.heading_blue:
                style_heading.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            # Per-level styling
            if h == 1:
                style_heading.font.size = Pt(22)
                style_heading.paragraph_format.space_after = Pt(2)
            elif h == 2:
                style_heading.font.size = Pt(17)
            elif h <= 4:
                style_heading.font.size = Pt(13)
            # Italics for small headings
            if h > 3:
                style_heading.font.italic = True

        # Styling for paragraphs
        style_paragraph = docx_doc.styles["Normal"]
        style_paragraph.font.size = Pt(self.style.body_pt)
        style_paragraph.paragraph_format.alignment = self.style._body_alignment()
        style_paragraph.paragraph_format.line_spacing = self.style.body_lines

        # Styling for captions
        if not self.style.heading_blue:
            style_caption = docx_doc.styles["Caption"]
            style_caption.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # Styling for codeblocks
        style_codeblock.font.name = self.style.font_code
        style_codeblock.paragraph_format.space_after = Pt(0)
        style_codeblock.paragraph_format.line_spacing = 1
        style_codeblock.paragraph_format.alignment = 0

        # TODO: new "Link" run styling, can be done

        # Use docx's vanilla save
        docx_doc.save(path)

